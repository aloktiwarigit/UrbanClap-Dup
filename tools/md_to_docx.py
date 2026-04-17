"""Convert the Hindi proposal markdown into a polished Word document.

Devanagari text is rendered with Nirmala UI (ships with Windows). Tables, headings,
bold/italic, and code blocks are handled. Designed to keep formatting close to the
markdown source without needing pandoc.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

SRC = Path(r"C:\Alok\Business Projects\Urbanclap-dup\docs\client-proposal-hindi.md")
DST = Path(r"C:\Alok\Business Projects\Urbanclap-dup\docs\client-proposal-hindi.docx")

DEVANAGARI_FONT = "Nirmala UI"
LATIN_FONT = "Calibri"
MONO_FONT = "Consolas"


def _set_run_fonts(run, size_pt: float, bold: bool = False, italic: bool = False,
                   mono: bool = False, color: RGBColor | None = None) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    latin = MONO_FONT if mono else LATIN_FONT
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), DEVANAGARI_FONT)
    rfonts.set(qn("w:eastAsia"), latin)


# Inline pattern: bold (**), italic (*), code (`)
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _add_inline(paragraph, text: str, base_size: float = 11) -> None:
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _set_run_fonts(run, base_size)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_fonts(run, base_size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_fonts(run, base_size - 1, mono=True,
                           color=RGBColor(0xC7, 0x25, 0x4E))
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_fonts(run, base_size, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_fonts(run, base_size)


def _add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_fonts(run, sizes.get(level, 11), bold=True,
                   color=RGBColor(0x1F, 0x3A, 0x5F))


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        _set_run_fonts(run, 10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tc_pr = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3A5F")
        tc_pr.append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_i, row in enumerate(rows, start=1):
        for c_i, cell_text in enumerate(row):
            if c_i >= len(header):
                continue
            cell = table.rows[r_i].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline(p, cell_text.strip(), base_size=10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _parse_table_block(lines: list[str], i: int) -> tuple[list[str], list[list[str]], int]:
    header_line = lines[i]
    sep_line = lines[i + 1]
    header = [c.strip() for c in header_line.strip().strip("|").split("|")]
    rows = []
    j = i + 2
    while j < len(lines) and lines[j].strip().startswith("|"):
        row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
        rows.append(row)
        j += 1
    return header, rows, j


def convert() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(11)

    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.4)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_buffer))
                _set_run_fonts(run, 9.5, mono=True,
                               color=RGBColor(0x33, 0x33, 0x33))
                shading = p._p.get_or_add_pPr()
                from docx.oxml import OxmlElement
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F2F2")
                shading.append(shd)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), 4)
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "BFBFBF")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif line.strip().startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\|[\s\-:|]+\|\s*$", lines[i + 1]):
            header, rows, next_i = _parse_table_block(lines, i)
            _add_table(doc, header, rows)
            i = next_i
            doc.add_paragraph()
            continue
        elif line.lstrip().startswith(("- ", "* ")):
            stripped = line.lstrip()
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, content)
        elif re.match(r"^\d+\.\s", line.lstrip()):
            content = re.sub(r"^\d+\.\s", "", line.lstrip())
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, content)
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_inline(p, line)

        i += 1

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"Wrote: {DST}")


if __name__ == "__main__":
    convert()
